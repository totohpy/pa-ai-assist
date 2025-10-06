import streamlit as st
from datetime import datetime
import json
import re
from openai import OpenAI
import os
import html
import io
import docx
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt

# --- Page Configuration ---
st.set_page_config(layout="wide", page_title="AI Plan Generator")
st.title("🤖 AI Plan Generator")
st.markdown("เครื่องมือช่วยสร้างแผนและแนวการตรวจสอบ พร้อมระบบ AI ช่วยร่างเนื้อหา")

# --- Custom CSS for Styling (same as before) ---
st.markdown("""
<style>
/* General expander button text */
div[data-testid="stExpander"] div[role="button"] p {
    font-size: 1.1rem;
}
.ai-expander .st-emotion-cache-ff2938 {
    background-color: #e7f3ff; border: 1px solid #007bff; border-radius: 0.5rem;
}
.ai-expander .st-emotion-cache-ff2938:hover {
    background-color: #d0e8ff;
}
.ai-expander .st-emotion-cache-ff2938 p {
    color: #004085; font-weight: bold;
}
.ai-button-container .stButton > button {
    background-color: #d4edda; color: #155724; border: 1px solid #c3e6cb;
    font-weight: bold; border-radius: 0.5rem; width: 100%;
}
.ai-button-container .stButton > button:hover {
    background-color: #c3e6cb; color: #155724; border-color: #b1dfbb;
}
</style>
""", unsafe_allow_html=True)

# --- State Initialization and Helper Functions (same as before) ---
def init_plan_state():
    ss = st.session_state
    if "plan_gen_data" not in ss:
        ss.plan_gen_data = {
            "general_info": {"office": "", "topic": "", "agency": "", "ministry": ""},
            "objectives": [],
            "estimates": {"cost": "", "effort": ""},
            "signatures": {
                "maker": {"name": "", "position": "", "date": None, "comment": ""},
                "reviewer": {"name": "", "position": "", "date": None, "comment": ""},
                "approver": {"name": "", "position": "", "date": None, "comment": ""},
            }
        }
    if "ui_feedback_message" not in ss:
        ss.ui_feedback_message = None
init_plan_state()

def add_objective():
    new_obj = {"id": f"obj_{len(st.session_state.plan_gen_data['objectives']) + 1}", "text": "", "issues": []}
    st.session_state.plan_gen_data["objectives"].append(new_obj)
    st.session_state.ui_feedback_message = None

def remove_objective(obj_index):
    st.session_state.plan_gen_data["objectives"].pop(obj_index)
    st.session_state.ui_feedback_message = None

def add_issue(obj_index, parent_issue_path=None):
    obj = st.session_state.plan_gen_data["objectives"][obj_index]
    target_container = obj
    if parent_issue_path:
        for index in parent_issue_path:
            target_container = target_container["issues"][index]
    new_issue = {
        "id": f"issue_{obj_index}_{len(target_container['issues']) + 1}", "text": "",
        "details": {"criteria": "", "info_needed": "", "source": "", "collection_method": "", "analysis_method": ""}, "issues": []
    }
    target_container["issues"].append(new_issue)
    st.session_state.ui_feedback_message = None

# --- AI Function (same as before) ---
def run_ai_for_field(obj_index, path, field_name):
    # ... This function remains unchanged from the previous version
    st.session_state.ui_feedback_message = None
    try:
        api_key = st.secrets["api_key"]
        client = OpenAI(api_key=api_key, base_url="https://api.opentyphoon.ai/v1")
        obj = st.session_state.plan_gen_data["objectives"][obj_index]
        target_issue = obj
        for index in path:
            target_issue = target_issue["issues"][index]
        context = f"เรื่องที่ตรวจสอบ: {st.session_state.plan_gen_data['general_info']['topic']}\n"
        context += f"วัตถุประสงค์: {obj.get('text', '')}\n"
        context += f"ประเด็นการตรวจสอบ: {target_issue.get('text', '')}\n"
        prompt_instruction = ""
        if field_name == "criteria":
            prompt_instruction = "จาก context ข้างต้น จงสร้างเฉพาะ 'เกณฑ์การตรวจสอบ' (Audit Criteria) ที่เหมาะสม"
        elif field_name == "info_needed":
            context += f"เกณฑ์การตรวจสอบ: {target_issue['details'].get('criteria', '')}\n"
            prompt_instruction = "จาก context ข้างต้น จงระบุ 'ข้อมูลที่ต้องการ' เพื่อสนับสนุนข้อตรวจพบ, สรุปผล, และข้อเสนอแนะ รวมถึงการตอบประเด็น, สภาพปัญหา, ผลกระทบ, และสาเหตุ"
        elif field_name == "source":
            context += f"เกณฑ์การตรวจสอบ: {target_issue['details'].get('criteria', '')}\n"
            context += f"ข้อมูลที่ต้องการ: {target_issue['details'].get('info_needed', '')}\n"
            prompt_instruction = "จาก context ข้างต้น จงระบุ 'แหล่งข้อมูล' ที่จะสามารถรวบรวมข้อมูลได้"
        elif field_name == "collection_method":
            context += f"เกณฑ์การตรวจสอบ: {target_issue['details'].get('criteria', '')}\n"
            context += f"ข้อมูลที่ต้องการ: {target_issue['details'].get('info_needed', '')}\n"
            context += f"แหล่งข้อมูล: {target_issue['details'].get('source', '')}\n"
            prompt_instruction = "จาก context ข้างต้น จงระบุ 'วิธีการรวบรวมหลักฐาน' เช่น การสุ่มตัวอย่าง, การตรวจสอบเอกสาร, การสัมภาษณ์, การสังเกตการณ์"
        elif field_name == "analysis_method":
            context += f"เกณฑ์การตรวจสอบ: {target_issue['details'].get('criteria', '')}\n"
            context += f"ข้อมูลที่ต้องการ: {target_issue['details'].get('info_needed', '')}\n"
            context += f"แหล่งข้อมูล: {target_issue['details'].get('source', '')}\n"
            context += f"วิธีการรวบรวมหลักฐาน: {target_issue['details'].get('collection_method', '')}\n"
            prompt_instruction = "จาก context ข้างต้น จงระบุ 'วิธีการวิเคราะห์หลักฐาน' ที่จะใช้ในการประมวลผล"
        full_prompt = f"คุณคือผู้เชี่ยวชาญด้านการตรวจสอบภาครัฐ\n{context}\n**คำสั่ง:**\n{prompt_instruction}\nตอบกลับเป็นข้อความธรรมดาในรูปแบบรายการ (bullet points) เท่านั้น"
        messages = [{"role": "user", "content": full_prompt}]
        response = client.chat.completions.create(
            model="typhoon-v2.1-12b-instruct", messages=messages, temperature=0.5
        )
        generated_text = response.choices[0].message.content.strip()
        cleaned_text = re.sub(r'^\s*[\*\-]\s*', '', generated_text, flags=re.MULTILINE)
        cleaned_text = cleaned_text.replace("**", "")
        if cleaned_text:
            target_issue['details'][field_name] = cleaned_text
            key_suffix = f"{obj_index}_{'_'.join(map(str, path))}"
            widget_key = f"{field_name}_{key_suffix}"
            st.session_state[widget_key] = cleaned_text
            st.session_state.ui_feedback_message = ("success", f"AI สร้าง '{field_name}' เรียบร้อยแล้ว")
        else:
            st.session_state.ui_feedback_message = ("error", f"AI ไม่สามารถสร้างเนื้อหาสำหรับ '{field_name}' ได้")
    except Exception as e:
        st.session_state.ui_feedback_message = ("error", f"เกิดข้อผิดพลาดในการเรียก AI: {e}")

# --- NEW: Function to generate DOCX report ---
def generate_docx_report(data):
    """Formats the plan data into a .docx file in memory."""
    doc = docx.Document()
    
    # Change orientation to landscape
    current_section = doc.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    current_section.orientation = WD_ORIENT.LANDSCAPE
    current_section.page_width = new_width
    current_section.page_height = new_height
    
    # Set font
    font = doc.styles['Normal'].font
    font.name = 'TH SarabunPSK'
    font.size = Pt(14)
    
    doc.add_heading('แผนและแนวการตรวจสอบ', level=1)
    
    info = data["general_info"]
    doc.add_paragraph(
        f"เรื่องที่ตรวจสอบ: {info.get('topic', 'N/A')}    "
        f"หน่วยงาน: {info.get('agency', 'N/A')}    "
        f"กระทรวง: {info.get('ministry', 'N/A')}"
    )
    doc.add_paragraph(f"สำนักงาน/จังหวัด/กลุ่ม: {info.get('office', 'N/A')}")
    
    for i, obj in enumerate(data["objectives"]):
        doc.add_paragraph().add_run(f"วัตถุประสงค์การตรวจสอบที่ {i+1}: {obj.get('text', '')}").bold = True
        
        for j, issue in enumerate(obj.get('issues', [])):
            doc.add_paragraph(f"ประเด็นการตรวจสอบที่ {i+1}.{j+1}: {issue.get('text', '')}")
            
            if not issue.get('issues'): # It's a leaf node with details
                details = issue.get('details', {})
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'เกณฑ์การตรวจสอบ'
                hdr_cells[1].text = 'ข้อมูลที่ต้องการ'
                hdr_cells[2].text = 'แหล่งข้อมูล'
                hdr_cells[3].text = 'วิธีการรวบรวมหลักฐาน'
                hdr_cells[4].text = 'วิธีการวิเคราะห์หลักฐาน'
                
                row_cells = table.add_row().cells
                row_cells[0].text = details.get('criteria', '')
                row_cells[1].text = details.get('info_needed', '')
                row_cells[2].text = details.get('source', '')
                row_cells[3].text = details.get('collection_method', '')
                row_cells[4].text = details.get('analysis_method', '')
                doc.add_paragraph() # Add space after table

    doc.add_heading('ประมาณการ', level=2)
    estimates = data["estimates"]
    doc.add_paragraph(f"ประมาณการค่าใช้จ่ายในการตรวจสอบ: {estimates.get('cost', '')}")
    doc.add_paragraph(f"ประมาณการคน/วันที่ใช้ในการตรวจสอบ: {estimates.get('effort', '')}")
    
    doc.add_heading('ผู้จัดทำและลงนาม', level=2)
    sigs = data["signatures"]
    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.style = 'Table Grid'
    hdr_cells = sig_table.rows[0].cells
    hdr_cells[0].text = 'ผู้จัดทำ'
    hdr_cells[1].text = 'ผู้สอบทาน'
    hdr_cells[2].text = 'ผู้อนุมัติ (รผต. / ผอ. สำนัก)'

    data_row = sig_table.add_row().cells
    for idx, role in enumerate(["maker", "reviewer", "approver"]):
        sig = sigs.get(role, {})
        date_str = sig.get('date').strftime('%Y-%m-%d') if sig.get('date') else ''
        cell_text = (
            f"ลงชื่อ: {sig.get('name', '')}\n"
            f"ตำแหน่ง: {sig.get('position', '')}\n"
            f"วันที่: {date_str}\n"
            f"ความเห็นเพิ่มเติม: {sig.get('comment', '')}"
        )
        data_row[idx].text = cell_text

    # Save to a memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- UI Rendering (Main part of the app) ---
# ... (The UI rendering code for sections 1, 2, 3 is the same as before) ...

if st.session_state.get("ui_feedback_message"):
    msg_type, msg_content = st.session_state.ui_feedback_message
    if msg_type == "success": st.success(msg_content)
    else: st.error(msg_content)
    st.session_state.ui_feedback_message = None

with st.form("general_info_form"):
    st.subheader("1. ข้อมูลทั่วไป")
    c1, c2 = st.columns(2)
    st.session_state.plan_gen_data["general_info"]["office"] = c1.text_input("สำนักงาน/จังหวัด/กลุ่ม", st.session_state.plan_gen_data["general_info"]["office"])
    st.session_state.plan_gen_data["general_info"]["topic"] = c1.text_input("เรื่องที่ตรวจสอบ", st.session_state.plan_gen_data["general_info"]["topic"])
    st.session_state.plan_gen_data["general_info"]["agency"] = c2.text_input("หน่วยงาน", st.session_state.plan_gen_data["general_info"]["agency"])
    st.session_state.plan_gen_data["general_info"]["ministry"] = c2.text_input("กระทรวง", st.session_state.plan_gen_data["general_info"]["ministry"])
    st.form_submit_button("บันทึกข้อมูลทั่วไป", use_container_width=True)

st.subheader("2. วัตถุประสงค์และประเด็นการตรวจสอบ")
for i, obj in enumerate(st.session_state.plan_gen_data["objectives"]):
    with st.container(border=True):
        c1, c2 = st.columns([5, 1])
        st.session_state.plan_gen_data["objectives"][i]['text'] = c1.text_area(f"วัตถุประสงค์ที่ {i+1}", obj.get('text', ''), key=f"obj_text_{i}")
        c2.button("🗑️ ลบ", key=f"del_obj_{i}", on_click=remove_objective, args=(i,), use_container_width=True)

        def display_issues(issues_list, obj_index, path):
            for j, issue in enumerate(issues_list):
                current_path = path + [j]
                prefix = ".".join([str(obj_index + 1)] + [str(p + 1) for p in current_path])
                key_suffix = f"{obj_index}_{'_'.join(map(str, current_path))}"

                with st.container():
                    st.markdown(f"<div style='margin-left: {len(current_path) * 20}px;'>", unsafe_allow_html=True)
                    target_issue = issue
                    target_issue['text'] = st.text_area(f"ประเด็นการตรวจสอบที่ {prefix}", value=target_issue.get('text', ''), key=f"issue_text_{key_suffix}")
                    if not target_issue.get('issues'):
                        st.markdown('<div class="ai-expander">', unsafe_allow_html=True)
                        with st.expander("เพิ่มรายละเอียดแนวการตรวจสอบ (ให้ AI ช่วย)"):
                            details = target_issue.get('details', {})
                            field_map = { "criteria": "เกณฑ์การตรวจสอบ", "info_needed": "ข้อมูลที่ต้องการ", "source": "แหล่งข้อมูล", "collection_method": "วิธีการรวบรวมหลักฐาน", "analysis_method": "วิธีการวิเคราะห์หลักฐาน" }
                            for field, label in field_map.items():
                                col1, col2 = st.columns([4, 1])
                                with col1:
                                    details[field] = st.text_area(label, value=details.get(field, ''), key=f"{field}_{key_suffix}")
                                with col2:
                                    st.markdown('<div class="ai-button-container">', unsafe_allow_html=True)
                                    st.button(f"✨สร้าง", key=f"ai_btn_{field}_{key_suffix}", on_click=run_ai_for_field, args=(i, current_path, field))
                                    st.markdown('</div>', unsafe_allow_html=True)
                        st.markdown('</div>', unsafe_allow_html=True)
                    st.button(f"➕ เพิ่มประเด็นย่อย (สำหรับ {prefix})", key=f"add_sub_issue_{key_suffix}", on_click=add_issue, args=(obj_index, current_path))
                    if target_issue.get('issues'): display_issues(target_issue['issues'], obj_index, current_path)
                    st.markdown("</div>", unsafe_allow_html=True)

        if obj.get('issues'): display_issues(obj['issues'], i, [])
        st.button("➕ เพิ่มประเด็นการตรวจสอบหลัก", key=f"add_issue_{i}", on_click=add_issue, args=(i, None))

st.button("➕ เพิ่มวัตถุประสงค์", on_click=add_objective, type="primary")

with st.form("estimates_signatures_form"):
    st.subheader("3. ประมาณการและผู้จัดทำ")
    st.session_state.plan_gen_data["estimates"]["cost"] = st.text_area("ประมาณการค่าใช้จ่ายในการตรวจสอบ", st.session_state.plan_gen_data["estimates"]["cost"])
    st.session_state.plan_gen_data["estimates"]["effort"] = st.text_area("ประมาณการคน/วันที่ใช้ในการตรวจสอบ", st.session_state.plan_gen_data["estimates"]["effort"])
    c1, c2, c3 = st.columns(3)
    sig_data = st.session_state.plan_gen_data["signatures"]
    with c1:
        st.markdown("**ผู้จัดทำ**")
        sig_data["maker"]["name"] = st.text_input("ลงชื่อ", value=sig_data["maker"].get("name", ""), key="maker_name")
        sig_data["maker"]["position"] = st.text_input("ตำแหน่ง", value=sig_data["maker"].get("position", ""), key="maker_pos")
        sig_data["maker"]["date"] = st.date_input("วันที่", value=sig_data["maker"].get("date"), key="maker_date")
        sig_data["maker"]["comment"] = st.text_area("ความเห็นเพิ่มเติม", value=sig_data["maker"].get("comment", ""), key="maker_comment")
    with c2:
        st.markdown("**ผู้สอบทาน**")
        sig_data["reviewer"]["name"] = st.text_input("ลงชื่อ", value=sig_data["reviewer"].get("name", ""), key="reviewer_name")
        sig_data["reviewer"]["position"] = st.text_input("ตำแหน่ง", value=sig_data["reviewer"].get("position", ""), key="reviewer_pos")
        sig_data["reviewer"]["date"] = st.date_input("วันที่", value=sig_data["reviewer"].get("date"), key="reviewer_date")
        sig_data["reviewer"]["comment"] = st.text_area("ความเห็นเพิ่มเติม", value=sig_data["reviewer"].get("comment", ""), key="reviewer_comment")
    with c3:
        st.markdown("**ผู้อนุมัติ (รผต. / ผอ. สำนัก)**")
        sig_data["approver"]["name"] = st.text_input("ลงชื่อ", value=sig_data["approver"].get("name", ""), key="approver_name")
        sig_data["approver"]["position"] = st.text_input("ตำแหน่ง", value=sig_data["approver"].get("position", ""), key="approver_pos")
        sig_data["approver"]["date"] = st.date_input("วันที่", value=sig_data["approver"].get("date"), key="approver_date")
        sig_data["approver"]["comment"] = st.text_area("ความเห็นเพิ่มเติม", value=sig_data["approver"].get("comment", ""), key="approver_comment")
    st.form_submit_button("บันทึกข้อมูลผู้จัดทำ", use_container_width=True)

# --- UPDATED: Section for saving and downloading the DOCX report ---
st.divider()
st.subheader("4. บันทึกและสร้างเอกสาร")

# Generate the docx file in memory
docx_buffer = generate_docx_report(st.session_state.plan_gen_data)

# Provide the download button for .docx
st.download_button(
    label="📂 ดาวน์โหลดแผนและแนวการตรวจ (.docx)",
    data=docx_buffer,
    file_name=f"audit_plan_{datetime.now().strftime('%Y-%m-%d')}.docx",
    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    use_container_width=True,
    type="primary"
)
