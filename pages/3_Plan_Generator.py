import streamlit as st
from datetime import datetime
import re
from openai import OpenAI
import os
import html
import io
import docx
from docx.enum.section import WD_ORIENT
from docx.shared import Pt
import base64
import json
import streamlit.components.v1 as components
from style import load_css  # 1. Import ฟังก์ชันจาก style.py

# --- Page Configuration ---
st.set_page_config(layout="wide", page_title="AI Plan Generator")

# --- Load CSS and Set Page State ---
load_css()  # 2. เรียกใช้ฟังก์ชันเพื่อโหลดสไตล์
st.session_state.current_page = "Plan Generator"  # 3. บอกแอปว่าตอนนี้อยู่ที่หน้าไหน

# --- Sidebar ---
with st.sidebar:
    st.title("เมนูหลัก")

    # --- สร้างปุ่มเมนู ---
    if st.button("หน้าหลัก (Home)", use_container_width=True, type="primary" if st.session_state.get("current_page") == "Home" else "secondary"):
        st.switch_page("Home.py")

    if st.button("Audit Design Assistant", use_container_width=True, type="primary" if st.session_state.get("current_page") == "Design Assistant" else "secondary"):
        st.switch_page("pages/2_Design_Assistant.py")

    if st.button("Audit Plan Generator", use_container_width=True, type="primary" if st.session_state.get("current_page") == "Plan Generator" else "secondary"):
        st.rerun() # อยู่หน้าตัวเอง ใช้ rerun

    if st.button("PA Assistant Chat", use_container_width=True, type="primary" if st.session_state.get("current_page") == "Chat" else "secondary"):
        st.switch_page("pages/4_PA_Assistant_Chat.py")

    # --- Footer ---
    st.markdown("""
        <div class="sidebar-footer">
            <p>
                <span style="color: grey;">By PAO1 </span><br>
                <span style="font-size: 16px; letter-spacing: 0.5px;">
                    <span style="color: red; font-weight: bold;">A</span>udit
                    <span style="color: red; font-weight: bold;">I</span>ntelligence
                    <span style="color: red; font-weight: bold;">T</span>eam
                </span>
            </p>
        </div>
    """, unsafe_allow_html=True)

# --- ❌ ไม่ต้องมี st.markdown("<style>...</style>") บล็อกใหญ่อีกต่อไป ---


# --- Main Page Content ---
st.title("🔮 Plan Generator")
st.markdown("เครื่องมือช่วยสร้างแผนและแนวการตรวจสอบ พร้อมระบบ AI ช่วยร่างเนื้อหา")

# --- State Initialization and Helper Functions ---
def init_plan_state():
    ss = st.session_state
    if "plan_gen_data" not in ss:
        ss.plan_gen_data = { "general_info": {"office": "", "topic": "", "agency": "", "ministry": ""}, "objectives": [], "estimates": {"cost": "", "effort": ""}, "signatures": { "maker": {"name": "", "position": "", "date": None, "comment": ""}, "reviewer": {"name": "", "position": "", "date": None, "comment": ""}, "approver": {"name": "", "position": "", "date": None, "comment": ""}, } }
    if "ui_feedback_message" not in ss:
        ss.ui_feedback_message = None
    # Load API Key from secrets
    if 'api_key_global' not in ss:
        try:
            ss['api_key_global'] = st.secrets["api_key"]
        except (KeyError, FileNotFoundError):
            ss['api_key_global'] = ""

init_plan_state()

def add_objective():
    new_obj = {"id": f"obj_{len(st.session_state.plan_gen_data['objectives']) + 1}", "text": "", "issues": []}
    st.session_state.plan_gen_data["objectives"].append(new_obj)
    st.session_state.ui_feedback_message = None

def remove_objective(obj_index):
    st.session_state.plan_gen_data["objectives"].pop(obj_index)
    st.session_state.ui_feedback_message = None

def add_issue(obj_index, parent_issue_path=None):
    obj = st.session_state.plan_gen_data["objectives"][obj_index]
    target_container = obj
    if parent_issue_path:
        for index in parent_issue_path:
            target_container = target_container["issues"][index]
    new_issue = { "id": f"issue_{obj_index}_{len(target_container['issues']) + 1}", "text": "", "details": {"criteria": "", "info_needed": "", "source": "", "collection_method": "", "analysis_method": ""}, "issues": [] }
    target_container["issues"].append(new_issue)
    st.session_state.ui_feedback_message = None


# --- AI Function ---
def run_ai_for_field(obj_index, path, field_name):
    st.session_state.ui_feedback_message = None
    try:
        api_key = st.session_state.get("api_key_global")
        if not api_key:
            st.session_state.ui_feedback_message = ("error", "ไม่พบ API Key กรุณาตั้งค่าใน Streamlit Cloud Secrets")
            return

        client = OpenAI(api_key=api_key, base_url="https://api.opentyphoon.ai/v1")
        obj = st.session_state.plan_gen_data["objectives"][obj_index]
        target_issue = obj
        for index in path:
            target_issue = target_issue["issues"][index]
        context = f"เรื่องที่ตรวจสอบ: {st.session_state.plan_gen_data['general_info']['topic']}\n"
        context += f"วัตถุประสงค์: {obj.get('text', '')}\n"
        context += f"ประเด็นการตรวจสอบ: {target_issue.get('text', '')}\n"
        prompt_instruction = ""
        if field_name == "criteria":
            prompt_instruction = "จาก context ข้างต้น จงสร้างเฉพาะ 'เกณฑ์การตรวจสอบ' (Audit Criteria) ที่เหมาะสม"
        elif field_name == "info_needed":
            context += f"เกณฑ์การตรวจสอบ: {target_issue['details'].get('criteria', '')}\n"
            prompt_instruction = "จาก context ข้างต้น จงระบุ 'ข้อมูลที่ต้องการ' เพื่อสนับสนุนข้อตรวจพบ, สรุปผล, และข้อเสนอแนะ รวมถึงการตอบประเด็น, สภาพปัญหา, ผลกระทบ, และสาเหตุ"
        elif field_name == "source":
            context += f"เกณฑ์การตรวจสอบ: {target_issue['details'].get('criteria', '')}\n"
            context += f"ข้อมูลที่ต้องการ: {target_issue['details'].get('info_needed', '')}\n"
            prompt_instruction = "จาก context ข้างต้น จงระบุ 'แหล่งข้อมูล' ที่จะสามารถรวบรวมข้อมูลได้"
        elif field_name == "collection_method":
            context += f"เกณฑ์การตรวจสอบ: {target_issue['details'].get('criteria', '')}\n"
            context += f"ข้อมูลที่ต้องการ: {target_issue['details'].get('info_needed', '')}\n"
            context += f"แหล่งข้อมูล: {target_issue['details'].get('source', '')}\n"
            prompt_instruction = "จาก context ข้างต้น จงระบุ 'วิธีการรวบรวมหลักฐาน' เช่น การสุ่มตัวอย่าง, การตรวจสอบเอกสาร, การสัมภาษณ์, การสังเกตการณ์"
        elif field_name == "analysis_method":
            context += f"เกณฑ์การตรวจสอบ: {target_issue['details'].get('criteria', '')}\n"
            context += f"ข้อมูลที่ต้องการ: {target_issue['details'].get('info_needed', '')}\n"
            context += f"แหล่งข้อมูล: {target_issue['details'].get('source', '')}\n"
            context += f"วิธีการรวบรวมหลักฐาน: {target_issue['details'].get('collection_method', '')}\n"
            prompt_instruction = "จาก context ข้างต้น จงระบุ 'วิธีการวิเคราะห์หลักฐาน' ที่จะใช้ในการประมวลผล"

        full_prompt = f"คุณคือผู้เชี่ยวชาญด้านการตรวจสอบภาครัฐ\n{context}\n**คำสั่ง:**\n{prompt_instruction}\nตอบกลับเป็นข้อความธรรมดาในรูปแบบรายการ (bullet points) เท่านั้น"
        messages = [{"role": "user", "content": full_prompt}]
        response = client.chat.completions.create(model="typhoon-v2.1-12b-instruct", messages=messages, temperature=0.5)
        generated_text = response.choices[0].message.content.strip()
        cleaned_text = generated_text.replace("**", "")

        if cleaned_text:
            target_issue['details'][field_name] = cleaned_text
            key_suffix = f"{obj_index}_{'_'.join(map(str, path))}"
            widget_key = f"{field_name}_{key_suffix}"
            st.session_state[widget_key] = cleaned_text
            st.session_state.ui_feedback_message = ("success", f"AI สร้าง '{field_name}' เรียบร้อยแล้ว")
        else:
            st.session_state.ui_feedback_message = ("error", f"AI ไม่สามารถสร้างเนื้อหาสำหรับ '{field_name}' ได้")
    except Exception as e:
        st.session_state.ui_feedback_message = ("error", f"เกิดข้อผิดพลาดในการเรียก AI: {e}")


# --- Functions to generate reports ---
@st.cache_data
def load_font_as_base64(font_path):
    try:
        with open(font_path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    except FileNotFoundError:
        st.warning(f"ไม่พบไฟล์ฟอนต์ที่ {font_path} จะใช้ฟอนต์มาตรฐานของเบราว์เซอร์แทน")
        return None

def generate_html_report(data):
    # This function remains unchanged from your original file
    sarabun_regular_b64 = load_font_as_base64("Sarabun-Regular.ttf")
    sarabun_bold_b64 = load_font_as_base64("Sarabun-Bold.ttf")
    sarabun_italic_b64 = load_font_as_base64("Sarabun-Italic.ttf")
    font_faces = ""
    if sarabun_regular_b64: font_faces += f"@font-face {{ font-family: 'Sarabun'; src: url(data:font/truetype;charset=utf-8;base64,{sarabun_regular_b64}) format('truetype'); font-weight: normal; font-style: normal; }}\n"
    if sarabun_bold_b64: font_faces += f"@font-face {{ font-family: 'Sarabun'; src: url(data:font/truetype;charset=utf-8;base64,{sarabun_bold_b64}) format('truetype'); font-weight: bold; font-style: normal; }}\n"
    if sarabun_italic_b64: font_faces += f"@font-face {{ font-family: 'Sarabun'; src: url(data:font/truetype;charset=utf-8;base64,{sarabun_italic_b64}) format('truetype'); font-weight: normal; font-style: italic; }}\n"
    def format_text(text):
        if not text: return ""
        return html.escape(text).replace('\n', '<br>')

    html_template = f"""
    <!DOCTYPE html><html lang="th"><head><meta charset="UTF-8"><title>แผนและแนวการตรวจสอบ</title>
    <style>
        {font_faces}
        body {{ font-family: 'Sarabun', sans-serif; font-size: 16px; margin: 1cm; }}
        h2 {{ text-align: center; font-weight: bold; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 1em; margin-bottom: 1em; }}
        th, td {{ border: 1px solid black; padding: 8px; text-align: left; vertical-align: top; }}
        .header-info p, .objective-info p {{ margin: 0; padding: 2px 0; }}
        .signature-table td {{ height: 120px; }}
        .print-button-container {{ text-align: center; margin: 15px 0; }}
        .print-button {{ padding: 10px 20px; font-size: 16px; cursor: pointer; border-radius: 5px; border: 1px solid #007bff; background-color: #007bff; color: white; font-family: 'Sarabun', sans-serif;}}
        thead th {{ position: -webkit-sticky; position: sticky; top: -1px; background-color: #f2f2f2; z-index: 2; }}
        @media print {{ .no-print {{ display: none; }} @page {{ size: A4 landscape; margin: 1.5cm; }} body {{ margin: 0; }} }}
    </style></head><body>
    <div class="print-button-container no-print"><button class="print-button" onclick="window.print()">🖨️ พิมพ์ / บันทึกเป็น PDF</button></div>
    <h2>แผนและแนวการตรวจสอบ</h2>
    <div class="header-info">
        <p><b>สำนักงานการตรวจเงินแผ่นดินภูมิภาคที่/สำนักตรวจเงินแผ่นดินจังหวัด:</b> {format_text(data['general_info'].get('office', '...........'))}
        <b>กลุ่มที่:</b> ........ <b>เรื่องที่ตรวจสอบ:</b> {format_text(data['general_info'].get('topic', '...........'))}</p>
        <p><b>หน่วยงาน:</b> {format_text(data['general_info'].get('agency', '...........'))}
        <b>กระทรวง:</b> {format_text(data['general_info'].get('ministry', '...........'))}</p>
    </div>
    {''.join([f'''<div class="objective-info"><p><b>วัตถุประสงค์การตรวจสอบที่ {i+1}:</b> {format_text(obj.get('text', ''))}</p></div>
    {''.join([f'''<table><thead><tr><th>เกณฑ์การตรวจสอบ</th><th>ข้อมูลที่ต้องการ</th><th>แหล่งข้อมูล</th><th>วิธีการรวบรวมหลักฐาน</th><th>วิธีการวิเคราะห์หลักฐาน</th></tr></thead>
    <tbody><tr><td colspan="5"><b>ประเด็นการตรวจสอบที่ {i+1}.{j+1}:</b> {format_text(issue.get('text', ''))}</td></tr>
    <tr><td>{format_text(issue['details'].get('criteria', ''))}</td><td>{format_text(issue['details'].get('info_needed', ''))}</td><td>{format_text(issue['details'].get('source', ''))}</td><td>{format_text(issue['details'].get('collection_method', ''))}</td><td>{format_text(issue['details'].get('analysis_method', ''))}</td></tr>
    </tbody></table>''' for j, issue in enumerate(obj.get('issues', [])) if not issue.get('issues')])}''' for i, obj in enumerate(data['objectives'])])}
    <p><b>ประมาณการค่าใช้จ่ายในการตรวจสอบ:</b><br>- {format_text(data['estimates'].get('cost', '..................'))}</p>
    <p><b>ประมาณการคน/วันที่ใช้ในการตรวจสอบ:</b><br>- {format_text(data['estimates'].get('effort', '..................'))}</p>
    <table class="signature-table"><thead><tr style="font-weight: bold; text-align: center;"><th>ผู้จัดทำ</th><th>ผู้สอบทาน</th><th>ผู้อนุมัติ (รผต. / ผอ. สำนัก)</th></tr></thead>
    <tbody><tr>
        <td><b>ลงชื่อ:</b> {format_text(data['signatures']['maker'].get('name', ''))}<br><b>ตำแหน่ง:</b> {format_text(data['signatures']['maker'].get('position', ''))}<br><b>วันที่:</b> {data['signatures']['maker'].get('date').strftime('%d/%m/%Y') if data['signatures']['maker'].get('date') else ''}<br><b>ความเห็นเพิ่มเติม:</b> {format_text(data['signatures']['maker'].get('comment', ''))}</td>
        <td><b>ลงชื่อ:</b> {format_text(data['signatures']['reviewer'].get('name', ''))}<br><b>ตำแหน่ง:</b> {format_text(data['signatures']['reviewer'].get('position', ''))}<br><b>วันที่:</b> {data['signatures']['reviewer'].get('date').strftime('%d/%m/%Y') if data['signatures']['reviewer'].get('date') else ''}<br><b>ความเห็นเพิ่มเติม:</b> {format_text(data['signatures']['reviewer'].get('comment', ''))}</td>
        <td><b>ลงชื่อ:</b> {format_text(data['signatures']['approver'].get('name', ''))}<br><b>ตำแหน่ง:</b> {format_text(data['signatures']['approver'].get('position', ''))}<br><b>วันที่:</b> {data['signatures']['approver'].get('date').strftime('%d/%m/%Y') if data['signatures']['approver'].get('date') else ''}<br><b>ความเห็นเพิ่มเติม:</b> {format_text(data['signatures']['approver'].get('comment', ''))}</td>
    </tr></tbody></table></body></html>
    """
    return html_template

def generate_docx_report(data):
    # This function remains unchanged from your original file
    doc = docx.Document()
    current_section = doc.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    current_section.orientation = WD_ORIENT.LANDSCAPE; current_section.page_width = new_width; current_section.page_height = new_height
    font = doc.styles['Normal'].font; font.name = 'TH SarabunPSK'; font.size = Pt(14)
    doc.add_heading('แผนและแนวการตรวจสอบ', level=1)
    info = data["general_info"]; doc.add_paragraph(f"เรื่องที่ตรวจสอบ: {info.get('topic', 'N/A')}    หน่วยงาน: {info.get('agency', 'N/A')}    กระทรวง: {info.get('ministry', 'N/A')}"); doc.add_paragraph(f"สำนักงาน/จังหวัด/กลุ่ม: {info.get('office', 'N/A')}")
    for i, obj in enumerate(data["objectives"]):
        doc.add_paragraph().add_run(f"วัตถุประสงค์การตรวจสอบที่ {i+1}: {obj.get('text', '')}").bold = True
        for j, issue in enumerate(obj.get('issues', [])):
            doc.add_paragraph(f"ประเด็นการตรวจสอบที่ {i+1}.{j+1}: {issue.get('text', '')}")
            if not issue.get('issues'):
                details = issue.get('details', {}); table = doc.add_table(rows=1, cols=5); table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells; hdr_cells[0].text = 'เกณฑ์การตรวจสอบ'; hdr_cells[1].text = 'ข้อมูลที่ต้องการ'; hdr_cells[2].text = 'แหล่งข้อมูล'; hdr_cells[3].text = 'วิธีการรวบรวมหลักฐาน'; hdr_cells[4].text = 'วิธีการวิเคราะห์หลักฐาน'
                row_cells = table.add_row().cells; row_cells[0].text = details.get('criteria', ''); row_cells[1].text = details.get('info_needed', ''); row_cells[2].text = details.get('source', ''); row_cells[3].text = details.get('collection_method', ''); row_cells[4].text = details.get('analysis_method', ''); doc.add_paragraph()
    doc.add_heading('ประมาณการ', level=2); estimates = data["estimates"]; doc.add_paragraph(f"ประมาณการค่าใช้จ่ายในการตรวจสอบ: {estimates.get('cost', '')}"); doc.add_paragraph(f"ประมาณการคน/วันที่ใช้ในการตรวจสอบ: {estimates.get('effort', '')}")
    doc.add_heading('ผู้จัดทำและลงนาม', level=2); sigs = data["signatures"]; sig_table = doc.add_table(rows=1, cols=3); sig_table.style = 'Table Grid'
    hdr_cells = sig_table.rows[0].cells; hdr_cells[0].text = 'ผู้จัดทำ'; hdr_cells[1].text = 'ผู้สอบทาน'; hdr_cells[2].text = 'ผู้อนุมัติ (รผต. / ผอ. สำนัก)'
    data_row = sig_table.add_row().cells
    for idx, role in enumerate(["maker", "reviewer", "approver"]):
        sig = sigs.get(role, {}); date_str = sig.get('date').strftime('%d/%m/%Y') if sig.get('date') else ''
        cell_text = f"ลงชื่อ: {sig.get('name', '')}\nตำแหน่ง: {sig.get('position', '')}\nวันที่: {date_str}\nความเห็นเพิ่มเติม: {sig.get('comment', '')}"; data_row[idx].text = cell_text
    buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer


# --- UI Rendering ---
if st.session_state.get("ui_feedback_message"):
    msg_type, msg_content = st.session_state.ui_feedback_message
    if msg_type == "success": st.success(msg_content)
    else: st.error(msg_content)
    st.session_state.ui_feedback_message = None

with st.form("general_info_form"):
    st.subheader("1. ข้อมูลทั่วไป")
    c1, c2 = st.columns(2)
    st.session_state.plan_gen_data["general_info"]["office"] = c1.text_input("สำนักงาน/จังหวัด/กลุ่ม", st.session_state.plan_gen_data["general_info"]["office"])
    st.session_state.plan_gen_data["general_info"]["topic"] = c1.text_input("เรื่องที่ตรวจสอบ", st.session_state.plan_gen_data["general_info"]["topic"])
    st.session_state.plan_gen_data["general_info"]["agency"] = c2.text_input("หน่วยงาน", st.session_state.plan_gen_data["general_info"]["agency"])
    st.session_state.plan_gen_data["general_info"]["ministry"] = c2.text_input("กระทรวง", st.session_state.plan_gen_data["general_info"]["ministry"])
    st.form_submit_button("บันทึกข้อมูลทั่วไป", use_container_width=True)

st.subheader("2. วัตถุประสงค์และประเด็นการตรวจสอบ")
for i, obj in enumerate(st.session_state.plan_gen_data["objectives"]):
    with st.container(border=True):
        c1, c2 = st.columns([5, 1])
        st.session_state.plan_gen_data["objectives"][i]['text'] = c1.text_area(f"วัตถุประสงค์ที่ {i+1}", obj.get('text', ''), key=f"obj_text_{i}")
        c2.button("🗑️ ลบ", key=f"del_obj_{i}", on_click=remove_objective, args=(i,), use_container_width=True)

        def display_issues(issues_list, obj_index, path):
            for j, issue in enumerate(issues_list):
                current_path = path + [j]
                prefix = ".".join([str(obj_index + 1)] + [str(p + 1) for p in current_path])
                key_suffix = f"{obj_index}_{'_'.join(map(str, current_path))}"
                with st.container():
                    st.markdown(f"<div style='margin-left: {len(current_path) * 20}px;'>", unsafe_allow_html=True)
                    target_issue = issue
                    target_issue['text'] = st.text_area(f"ประเด็นการตรวจสอบที่ {prefix}", value=target_issue.get('text', ''), key=f"issue_text_{key_suffix}")
                    if not target_issue.get('issues'):
                        with st.expander("เพิ่มรายละเอียดแนวการตรวจสอบ (🪄 AI ช่วย)"):
                            details = target_issue.get('details', {})
                            field_map = { "criteria": "เกณฑ์การตรวจสอบ", "info_needed": "ข้อมูลที่ต้องการ", "source": "แหล่งข้อมูล", "collection_method": "วิธีการรวบรวมหลักฐาน", "analysis_method": "วิธีการวิเคราะห์หลักฐาน" }
                            for field, label in field_map.items():
                                col1, col2 = st.columns([4, 1])
                                with col1:
                                    details[field] = st.text_area(label, value=details.get(field, ''), key=f"{field}_{key_suffix}")
                                with col2:
                                    st.button(f"✨สร้าง", key=f"ai_btn_{field}_{key_suffix}", on_click=run_ai_for_field, args=(i, current_path, field), use_container_width=True)
                    st.button(f"➕ เพิ่มประเด็นย่อย (สำหรับ {prefix})", key=f"add_sub_issue_{key_suffix}", on_click=add_issue, args=(obj_index, current_path))
                    if target_issue.get('issues'):
                        display_issues(target_issue['issues'], obj_index, current_path)
                    st.markdown("</div>", unsafe_allow_html=True)

        if obj.get('issues'):
            display_issues(obj['issues'], i, [])
        st.button("➕ เพิ่มประเด็นการตรวจสอบหลัก", key=f"add_issue_{i}", on_click=add_issue, args=(i, None))

st.button("➕ เพิ่มวัตถุประสงค์", on_click=add_objective, type="primary")

with st.form("estimates_signatures_form"):
    st.subheader("3. ประมาณการและผู้จัดทำ")
    st.session_state.plan_gen_data["estimates"]["cost"] = st.text_area("ประมาณการค่าใช้จ่ายในการตรวจสอบ", st.session_state.plan_gen_data["estimates"]["cost"])
    st.session_state.plan_gen_data["estimates"]["effort"] = st.text_area("ประมาณการคน/วันที่ใช้ในการตรวจสอบ", st.session_state.plan_gen_data["estimates"]["effort"])
    c1, c2, c3 = st.columns(3)
    sig_data = st.session_state.plan_gen_data["signatures"]
    with c1:
        st.markdown("**ผู้จัดทำ**")
        sig_data["maker"]["name"] = st.text_input("ลงชื่อ", value=sig_data["maker"].get("name", ""), key="maker_name")
        sig_data["maker"]["position"] = st.text_input("ตำแหน่ง", value=sig_data["maker"].get("position", ""), key="maker_pos")
        sig_data["maker"]["date"] = st.date_input("วันที่", value=sig_data["maker"].get("date"), key="maker_date")
        sig_data["maker"]["comment"] = st.text_area("ความเห็นเพิ่มเติม", value=sig_data["maker"].get("comment", ""), key="maker_comment")
    with c2:
        st.markdown("**ผู้สอบทาน**")
        sig_data["reviewer"]["name"] = st.text_input("ลงชื่อ", value=sig_data["reviewer"].get("name", ""), key="reviewer_name")
        sig_data["reviewer"]["position"] = st.text_input("ตำแหน่ง", value=sig_data["reviewer"].get("position", ""), key="reviewer_pos")
        sig_data["reviewer"]["date"] = st.date_input("วันที่", value=sig_data["reviewer"].get("date"), key="reviewer_date")
        sig_data["reviewer"]["comment"] = st.text_area("ความเห็นเพิ่มเติม", value=sig_data["reviewer"].get("comment", ""), key="reviewer_comment")
    with c3:
        st.markdown("**ผู้อนุมัติ (รผต. / ผอ. สำนัก)**")
        sig_data["approver"]["name"] = st.text_input("ลงชื่อ", value=sig_data["approver"].get("name", ""), key="approver_name")
        sig_data["approver"]["position"] = st.text_input("ตำแหน่ง", value=sig_data["approver"].get("position", ""), key="approver_pos")
        sig_data["approver"]["date"] = st.date_input("วันที่", value=sig_data["approver"].get("date"), key="approver_date")
        sig_data["approver"]["comment"] = st.text_area("ความเห็นเพิ่มเติม", value=sig_data["approver"].get("comment", ""), key="approver_comment")
    st.form_submit_button("บันทึกข้อมูลผู้จัดทำ", use_container_width=True)


# --- Section for direct HTML preview ---
st.divider()

with st.container(border=True):
    st.subheader("4. แสดงผลและส่งออกเอกสาร")
    html_report = generate_html_report(st.session_state.plan_gen_data)
    components.html(html_report, height=800, scrolling=True)

    st.markdown("---")

    st.markdown("##### หรือดาวน์โหลดเป็นไฟล์ Word")
    docx_buffer = generate_docx_report(st.session_state.plan_gen_data)
    st.download_button(
        label="📂 ดาวน์โหลดเป็นไฟล์ Word (.docx)",
        data=docx_buffer,
        file_name=f"audit_plan_{datetime.now().strftime('%Y%m%d')}.docx",
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        use_container_width=True
    )

